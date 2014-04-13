"""
knowledge: transcriber.py

Module for transcribing and organizing information and
images into word and PDF documents.
"""
import time
import sys
import os
import docx


# Directory to output documents
OUTPUT_DIR = 'studysheets/'


def generate_doc(content, source):
    """
    Public: (List, String) -> None

    Creates a word document, formatted based on source.
    """
    _check_dir_exists()
    if source in WORD_FORMATTING:
        WORD_FORMATTING[source](content, source)
    else:
        sys.exit('Formatting for %s is missing' % source)


def _check_dir_exists():
    """
    Private: None -> None

    Checks if OUTPUT_DIR exists, if not create it.
    """
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        print "Creating new folder: %s" % OUTPUT_DIR


def _save_document(document, source):
    """
    Private: (String) -> None

    Saves document and attaches time to filename.
    """
    timenow = time.strftime("-%b-%d-%Y-%H:%M")
    docname = OUTPUT_DIR + source + timenow + '.docx'
    document.save(docname)
    print "Saved document in: %s" % docname


def _format_wikipedia_doc(content, source):
    """
    Private: (List) -> None

    Generates document with custom wikipedia formatting.
    """
    document = docx.Document()
    for entry in content:
        document.add_heading(entry['title'], 1)
        document.add_paragraph(entry['intro'])
    _save_document(document, source)


def _format_wolfram_doc(content, source):
    """
    Private: (List) -> None

    Generates document with custom formatting for wolfram.
    """
    pass


WORD_FORMATTING = {
    'wikipedia' : _format_wikipedia_doc,
    'wolfram'   : _format_wolfram_doc
}
